"""Report generation for normalized ThermoAnalyzer result records."""

from __future__ import annotations

import csv
import io
import datetime
from typing import Any, Optional, Union

from core.batch_runner import normalize_batch_summary_rows, summarize_batch_outcomes
from core.processing_schema import ensure_processing_payload
from core.result_serialization import flatten_result_records, partition_results_by_status, split_valid_results
from core.scientific_sections import condense_warning_limitations, scientific_context_to_report_sections
from utils.reference_data import find_nearest_reference

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError as exc:  # pragma: no cover
    raise ImportError(
        "python-docx is required for DOCX report generation. Install it with: pip install python-docx"
    ) from exc


_MAIN_METADATA_KEYS = (
    "sample_name",
    "sample_mass",
    "heating_rate",
    "instrument",
    "vendor",
    "atmosphere",
    "display_name",
    "file_name",
    "import_confidence",
    "import_confidence_label",
    "inferred_analysis_type",
)

_APPENDIX_METADATA_KEYWORDS = (
    "hash",
    "delimiter",
    "decimal",
    "header",
    "row",
    "start",
    "import",
    "warning",
    "heuristic",
    "inferred",
    "parser",
    "dialect",
    "encoding",
)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_key_value_table(doc: Document, data: dict) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Parameter"
    hdr_cells[1].text = "Value"
    for cell in hdr_cells:
        _set_cell_bg(cell, "4472C4")
        for para in cell.paragraphs:
            run = para.runs[0] if para.runs else para.add_run(cell.text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for key, value in data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(value)


def _add_results_table(doc: Document, headers: list[str], rows: list[list]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        _set_cell_bg(hdr_cells[i], "4472C4")
        for para in hdr_cells[i].paragraphs:
            run = para.runs[0] if para.runs else para.add_run(hdr_cells[i].text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value)
            if row_idx % 2 == 1:
                _set_cell_bg(row_cells[i], "DCE6F1")


def _format_value(value):
    if value is None:
        return "N/A"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, list):
        return "; ".join(str(item) for item in value) if value else "N/A"
    if isinstance(value, dict):
        return ", ".join(f"{key}={_format_value(item)}" for key, item in value.items()) if value else "N/A"
    try:
        return f"{float(value):.4f}"
    except (TypeError, ValueError):
        return str(value)


def _record_title(record: dict) -> str:
    dataset_key = record.get("dataset_key")
    if dataset_key:
        return f"{record['analysis_type']} - {dataset_key}"
    return record["analysis_type"]


def _record_headers(record: dict) -> list[str]:
    rows = record.get("rows") or []
    if not rows:
        return []
    first_row = rows[0]
    return list(first_row.keys())


def _humanize_key(key: str) -> str:
    replacements = {"id": "ID", "utc": "UTC", "dsc": "DSC", "tga": "TGA", "dtg": "DTG"}
    words = str(key).replace("_", " ").split()
    return " ".join(replacements.get(word.lower(), word.capitalize()) for word in words)


def _table_payload(payload: dict | None) -> dict[str, str]:
    table_rows: dict[str, str] = {}
    for key, value in (payload or {}).items():
        if value in (None, "", [], {}):
            continue
        table_rows[_humanize_key(key)] = _format_value(value)
    return table_rows


def _safe_float(value: Any) -> float | None:
    try:
        if value in (None, ""):
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _format_number(value: Any, *, digits: int = 2) -> str:
    numeric = _safe_float(value)
    if numeric is None:
        return "N/A"
    return f"{numeric:.{digits}f}"


def _dataset_label(dataset_key: str | None, datasets: dict) -> str:
    if not dataset_key:
        return "N/A"
    dataset = datasets.get(dataset_key)
    if dataset is None:
        return dataset_key
    metadata = getattr(dataset, "metadata", {}) or {}
    return str(metadata.get("display_name") or metadata.get("sample_name") or metadata.get("file_name") or dataset_key)


def _main_conditions_payload(dataset_key: str, dataset) -> dict[str, str]:
    metadata = getattr(dataset, "metadata", {}) or {}
    payload = {
        "Dataset": _dataset_label(dataset_key, {dataset_key: dataset}),
        "Dataset Key": dataset_key,
        "Sample Name": metadata.get("sample_name"),
        "Sample Mass": metadata.get("sample_mass"),
        "Heating Rate": metadata.get("heating_rate"),
        "Instrument": metadata.get("instrument"),
        "Vendor": metadata.get("vendor"),
        "Atmosphere": metadata.get("atmosphere"),
        "Source File": metadata.get("file_name") or metadata.get("display_name"),
        "Import Confidence": metadata.get("import_confidence_label") or metadata.get("import_confidence"),
    }
    if payload.get("Import Confidence") in (None, "") and metadata.get("inferred_analysis_type"):
        payload["Import Confidence"] = f"Inferred type: {metadata.get('inferred_analysis_type')}"
    return _table_payload(payload)


def _appendix_dataset_metadata(metadata: dict | None) -> dict[str, str]:
    metadata = metadata or {}
    appendix_payload = {}
    for key, value in metadata.items():
        if value in (None, "", [], {}):
            continue
        lower = str(key).lower()
        if key in _MAIN_METADATA_KEYS:
            continue
        if any(token in lower for token in _APPENDIX_METADATA_KEYWORDS):
            appendix_payload[key] = value
    return _table_payload(appendix_payload)


def _record_key_results(record: dict) -> dict[str, str]:
    summary = record.get("summary") or {}
    analysis_type = str(record.get("analysis_type") or "").upper()
    if analysis_type == "TGA":
        keep = ("step_count", "total_mass_loss_percent", "residue_percent", "sample_name", "sample_mass", "heating_rate")
    elif analysis_type == "DSC":
        keep = (
            "peak_count",
            "glass_transition_count",
            "tg_midpoint",
            "tg_onset",
            "tg_endset",
            "delta_cp",
            "sample_name",
            "sample_mass",
            "heating_rate",
        )
    else:
        keep = tuple(summary.keys())
    return _table_payload({key: summary.get(key) for key in keep})


def _record_metric_snapshot(record: dict) -> str:
    summary = record.get("summary") or {}
    analysis_type = str(record.get("analysis_type") or "").upper()
    if analysis_type == "TGA":
        return ", ".join(
            [
                f"total mass loss {_format_number(summary.get('total_mass_loss_percent'))}%",
                f"residue {_format_number(summary.get('residue_percent'))}%",
                f"step count {_format_value(summary.get('step_count'))}",
            ]
        )
    if analysis_type == "DSC":
        snapshot = [f"peak count {_format_value(summary.get('peak_count'))}"]
        if summary.get("tg_midpoint") is not None:
            snapshot.append(f"Tg midpoint {_format_number(summary.get('tg_midpoint'))} °C")
        return ", ".join(snapshot)

    parts = []
    for key, value in summary.items():
        if key in {"sample_name", "sample_mass", "heating_rate"}:
            continue
        parts.append(f"{_humanize_key(key)} {_format_value(value)}")
        if len(parts) >= 3:
            break
    return ", ".join(parts) if parts else "Key metrics not available"


def _record_confidence_note(record: dict) -> str:
    grouped = condense_warning_limitations(
        record.get("scientific_context"),
        validation=record.get("validation"),
        max_data_warnings=2,
        max_method_limits=2,
    )
    data_warnings = grouped.get("data_completeness_warnings") or []
    method_limits = grouped.get("methodological_limitations") or []
    if data_warnings:
        return data_warnings[0]
    if method_limits:
        return method_limits[0]

    status = ((record.get("validation") or {}).get("status") or "").lower()
    if status == "pass":
        return "Validation checks did not flag material data-quality issues."
    if status == "warn":
        return "Validation generated warnings; interpret with moderate caution."
    if status == "fail":
        return "Validation failed; interpretation is not suitable for definitive conclusions."
    return "No explicit confidence constraints were recorded."


def _record_compact_rows(record: dict, *, max_rows: int = 5) -> tuple[list[str], list[list[str]]] | None:
    rows = record.get("rows") or []
    if not rows:
        return None
    if str(record.get("analysis_type") or "").upper() == "TGA":
        return None

    headers = _record_headers(record)
    if not headers:
        return None

    if len(rows) <= max_rows:
        return (
            headers,
            [[_format_value(row.get(header)) for header in headers] for row in rows],
        )

    top_headers = headers[: min(5, len(headers))]
    top_rows = [[_format_value(row.get(header)) for header in top_headers] for row in rows[:max_rows]]
    return top_headers, top_rows


def _tga_major_events(record: dict, *, limit: int = 3) -> list[list[str]]:
    if str(record.get("analysis_type") or "").upper() != "TGA":
        return []

    scored: list[tuple[float, dict]] = []
    for row in (record.get("rows") or []):
        if not isinstance(row, dict):
            continue
        loss = _safe_float(row.get("mass_loss_percent"))
        scored.append((loss if loss is not None else float("-inf"), row))

    scored.sort(key=lambda item: item[0], reverse=True)
    top = [row for _, row in scored[:limit] if row]

    output = []
    for idx, row in enumerate(top, start=1):
        output.append(
            [
                f"Event {idx}",
                _format_number(row.get("midpoint_temperature")),
                _format_number(row.get("mass_loss_percent")),
                _format_number(row.get("residual_percent")),
            ]
        )
    return output


def _record_full_rows(record: dict) -> tuple[list[str], list[list[str]]] | None:
    headers = _record_headers(record)
    if not headers:
        return None
    rows = [[_format_value(row.get(header)) for header in headers] for row in (record.get("rows") or [])]
    if not rows:
        return None
    return headers, rows


def _select_record_for_dataset(records: list[dict], dataset_key: str, analysis_type: str | None) -> dict | None:
    candidates = [record for record in records if record.get("dataset_key") == dataset_key]
    if not candidates:
        return None
    if analysis_type:
        normalized = analysis_type.upper()
        filtered = [record for record in candidates if str(record.get("analysis_type") or "").upper() == normalized]
        if filtered:
            stable = [record for record in filtered if record.get("status") == "stable"]
            return stable[0] if stable else filtered[0]
    stable = [record for record in candidates if record.get("status") == "stable"]
    return stable[0] if stable else candidates[0]


def _comparison_metadata_note(selected: list[str], datasets: dict) -> str:
    missing: list[str] = []
    for key, label in (("heating_rate", "heating-rate"), ("atmosphere", "atmosphere")):
        if any(not ((getattr(datasets.get(dataset_key), "metadata", {}) or {}).get(key)) for dataset_key in selected):
            missing.append(label)
    if not missing:
        return ""
    if len(missing) == 1:
        return f"Because {missing[0]} metadata are incomplete, this interpretation should be treated as comparative rather than definitive."
    return f"Because {', '.join(missing[:-1])}, and {missing[-1]} metadata are incomplete, this interpretation should be treated as comparative rather than definitive."


def _build_tga_comparison_interpretation(metrics: list[dict[str, Any]], metadata_note: str) -> str:
    valid = [item for item in metrics if item.get("mass_loss") is not None and item.get("residue") is not None and item.get("step_count") is not None]
    if len(valid) >= 2:
        left = valid[0]
        right = valid[1]
        left_loss = float(left["mass_loss"])
        right_loss = float(right["mass_loss"])
        left_residue = float(left["residue"])
        right_residue = float(right["residue"])
        left_steps = int(left["step_count"])
        right_steps = int(right["step_count"])

        step_phrase = "fewer" if right_steps < left_steps else "more" if right_steps > left_steps else "a similar number of"
        loss_phrase = "higher" if right_loss > left_loss else "lower" if right_loss < left_loss else "similar"
        residue_phrase = "lower" if right_residue < left_residue else "higher" if right_residue > left_residue else "similar"

        if right_loss > left_loss and right_residue < left_residue:
            implication = "suggesting a more complete mass-loss process under the recorded conditions"
        elif right_loss < left_loss and right_residue > left_residue:
            implication = "suggesting comparatively lower decomposition extent under the recorded conditions"
        else:
            implication = "indicating a materially different decomposition profile under the recorded conditions"

        text = (
            f"Compared with {left['dataset']}, {right['dataset']} shows {step_phrase} decomposition steps, "
            f"{loss_phrase} total mass loss ({right_loss:.2f}% vs {left_loss:.2f}%), and "
            f"{residue_phrase} final residue ({right_residue:.2f}% vs {left_residue:.2f}%), {implication}."
        )
    elif valid:
        only = valid[0]
        text = (
            f"{only['dataset']} reports total mass loss of {float(only['mass_loss']):.2f}%, "
            f"final residue of {float(only['residue']):.2f}%, and {int(only['step_count'])} resolved decomposition steps."
        )
    else:
        text = "Comparison metrics could not be resolved from the available TGA analysis results."

    if metadata_note:
        text = f"{text} {metadata_note}"
    return text


def _build_generic_comparison_interpretation(analysis_type: str, metrics: list[dict[str, Any]], metadata_note: str) -> str:
    if len(metrics) >= 2:
        text = (
            f"The {analysis_type} comparison shows measurable differences across the selected datasets. "
            "Review key metrics together with method context before drawing definitive mechanistic conclusions."
        )
    elif metrics:
        text = f"The selected {analysis_type} dataset was summarized successfully for comparison reporting."
    else:
        text = f"No {analysis_type} metrics were available for comparison interpretation."

    if metadata_note:
        text = f"{text} {metadata_note}"
    return text


def _build_comparison_payload(comparison_workspace: dict | None, datasets: dict, records: list[dict]) -> dict[str, Any] | None:
    comparison_workspace = comparison_workspace or {}
    selected = comparison_workspace.get("selected_datasets") or []
    if not selected:
        return None

    analysis_type = str(comparison_workspace.get("analysis_type") or "N/A")
    normalized_analysis = analysis_type.upper()

    overview = {
        "Analysis Type": analysis_type,
        "Compared Datasets": ", ".join(_dataset_label(dataset_key, datasets) for dataset_key in selected),
        "Saved Figure": comparison_workspace.get("figure_key") or "Not recorded",
    }

    metric_headers: list[str]
    metric_rows: list[list[str]] = []
    metric_records: list[dict[str, Any]] = []

    if normalized_analysis == "TGA":
        metric_headers = ["Dataset", "Total Mass Loss (%)", "Final Residue (%)", "Step Count"]
        for dataset_key in selected:
            dataset_name = _dataset_label(dataset_key, datasets)
            record = _select_record_for_dataset(records, dataset_key, normalized_analysis)
            summary = (record or {}).get("summary") or {}
            mass_loss = _safe_float(summary.get("total_mass_loss_percent"))
            residue = _safe_float(summary.get("residue_percent"))
            step_count_raw = summary.get("step_count")
            step_count = int(step_count_raw) if isinstance(step_count_raw, (int, float)) else None
            metric_records.append(
                {
                    "dataset": dataset_name,
                    "mass_loss": mass_loss,
                    "residue": residue,
                    "step_count": step_count,
                }
            )
            metric_rows.append(
                [
                    dataset_name,
                    _format_number(mass_loss),
                    _format_number(residue),
                    _format_value(step_count),
                ]
            )
        interpretation = _build_tga_comparison_interpretation(metric_records, _comparison_metadata_note(selected, datasets))
    else:
        metric_headers = ["Dataset", "Primary Metrics"]
        for dataset_key in selected:
            dataset_name = _dataset_label(dataset_key, datasets)
            record = _select_record_for_dataset(records, dataset_key, normalized_analysis)
            metric = _record_metric_snapshot(record or {}) if record else "No matching analysis result"
            metric_rows.append([dataset_name, metric])
            metric_records.append({"dataset": dataset_name})
        interpretation = _build_generic_comparison_interpretation(analysis_type, metric_records, _comparison_metadata_note(selected, datasets))

    appendix_overview = _table_payload(
        {
            "saved_at": comparison_workspace.get("saved_at"),
            "notes": comparison_workspace.get("notes"),
            "batch_run_id": comparison_workspace.get("batch_run_id"),
            "batch_template_label": comparison_workspace.get("batch_template_label"),
            "batch_template_id": comparison_workspace.get("batch_template_id"),
            "batch_completed_at": comparison_workspace.get("batch_completed_at"),
        }
    )

    return {
        "overview": _table_payload(overview),
        "metric_headers": metric_headers,
        "metric_rows": metric_rows,
        "interpretation": interpretation,
        "appendix_overview": appendix_overview,
        "appendix_batch_rows": _comparison_batch_rows(comparison_workspace),
    }


def _build_executive_summary_rows(records: list[dict], datasets: dict, comparison_payload: dict[str, Any] | None) -> list[list[str]]:
    rows: list[list[str]] = []

    for record in records:
        dataset_name = _dataset_label(record.get("dataset_key"), datasets)
        analysis_type = str(record.get("analysis_type") or "Analysis")
        status = str(record.get("status") or "unknown")
        interpretation_sections = scientific_context_to_report_sections(record.get("scientific_context"))
        interpretation_line = "No interpretation statement recorded."
        for title, payload in interpretation_sections:
            if title == "Scientific Interpretation" and isinstance(payload, dict) and payload:
                interpretation_line = str(next(iter(payload.values())))
                break
        rows.append(
            [
                dataset_name,
                f"{analysis_type} ({status})",
                _record_metric_snapshot(record),
                interpretation_line,
                _record_confidence_note(record),
            ]
        )

    if comparison_payload:
        rows.append(
            [
                "Comparison Set",
                "Cross-dataset comparison",
                "; ".join(
                    [
                        f"{comparison_payload['metric_headers'][idx]} captured"
                        for idx in range(1, len(comparison_payload.get("metric_headers") or []))
                    ]
                )
                or "Comparison metrics summarized",
                comparison_payload.get("interpretation") or "Comparison interpretation unavailable",
                "Comparative interpretation depends on metadata completeness and method parity.",
            ]
        )

    return rows


def _processing_step(processing: dict | None, key: str) -> dict:
    processing = processing or {}
    nested = processing.get("signal_pipeline") or {}
    if key in nested and isinstance(nested[key], dict):
        return nested[key]

    nested = processing.get("analysis_steps") or {}
    if key in nested and isinstance(nested[key], dict):
        return nested[key]

    value = processing.get(key)
    return value if isinstance(value, dict) else {}


def _format_processing_step(payload: dict | None) -> str:
    payload = payload or {}
    if not payload:
        return "Not recorded"

    method = payload.get("method")
    details = []
    for key, value in payload.items():
        if key == "method" or value in (None, "", [], {}):
            continue
        details.append(f"{_humanize_key(key)}={_format_value(value)}")

    if method and details:
        return f"{method} ({'; '.join(details)})"
    if method:
        return str(method)
    if details:
        return "; ".join(details)
    return "Recorded"


def _reference_visibility(record: dict) -> str:
    analysis_type = record.get("analysis_type")
    if analysis_type not in {"DSC", "TGA"}:
        return "Not applicable"

    candidate_temperature = None
    rows = record.get("rows") or []
    summary = record.get("summary") or {}
    if analysis_type == "DSC":
        if rows:
            candidate_temperature = rows[0].get("peak_temperature")
        if candidate_temperature in (None, ""):
            candidate_temperature = summary.get("tg_midpoint")
    elif analysis_type == "TGA":
        if rows:
            candidate_temperature = rows[0].get("midpoint_temperature") or rows[0].get("onset_temperature")

    if candidate_temperature in (None, ""):
        return "No reference candidate recorded"

    try:
        candidate_temperature = float(candidate_temperature)
    except (TypeError, ValueError):
        return "Reference candidate could not be parsed"

    reference = find_nearest_reference(candidate_temperature, analysis_type=analysis_type)
    if reference is None:
        return f"No close reference match within 15.0 °C (candidate {candidate_temperature:.1f} °C)"

    delta = candidate_temperature - reference.temperature_c
    standard = f"; {reference.standard}" if reference.standard else ""
    return f"{reference.name} ({reference.temperature_c:.1f} °C, ΔT {delta:+.1f} °C{standard})"


def _domain_method_summary(record: dict) -> dict[str, str] | None:
    analysis_type = record.get("analysis_type")
    if analysis_type not in {"DSC", "TGA"}:
        return None

    processing = ensure_processing_payload(record.get("processing"), analysis_type=analysis_type) if record.get("processing") else {}
    method_context = processing.get("method_context") or {}

    if analysis_type == "DSC":
        summary = {
            "Template": processing.get("workflow_template_label") or processing.get("workflow_template") or "Not recorded",
            "Sign Convention": method_context.get("sign_convention_label") or processing.get("sign_convention") or "Not recorded",
            "Smoothing": _format_processing_step(_processing_step(processing, "smoothing")),
            "Baseline": _format_processing_step(_processing_step(processing, "baseline")),
            "Peak Analysis Context": _format_processing_step(_processing_step(processing, "peak_detection")),
            "Glass Transition Context": _format_processing_step(_processing_step(processing, "glass_transition")),
            "Reference Check": _reference_visibility(record),
        }
        return _table_payload(summary)

    summary = {
        "Template": processing.get("workflow_template_label") or processing.get("workflow_template") or "Not recorded",
        "Declared Unit Mode": method_context.get("tga_unit_mode_label") or "Not recorded",
        "Resolved Unit Mode": method_context.get("tga_unit_mode_resolved_label") or "Not recorded",
        "Smoothing": _format_processing_step(_processing_step(processing, "smoothing")),
        "Step Analysis Context": _format_processing_step(_processing_step(processing, "step_detection")),
        "Reference Check": _reference_visibility(record),
    }
    return _table_payload(summary)


def _generic_method_summary(processing: dict | None) -> dict[str, str]:
    processing = processing or {}
    return _table_payload(
        {
            "analysis_type": processing.get("analysis_type"),
            "workflow_template": processing.get("workflow_template"),
            "method": processing.get("method"),
        }
    )


def _processing_sections(processing: dict | None) -> list[tuple[str, dict[str, str]]]:
    processing = processing or {}
    if not processing:
        return []

    sections: list[tuple[str, dict[str, str]]] = []
    method_context = _table_payload(processing.get("method_context"))
    if method_context:
        sections.append(("Method Context", method_context))

    signal_pipeline = processing.get("signal_pipeline")
    if not isinstance(signal_pipeline, dict) or not signal_pipeline:
        signal_pipeline = {
            key: processing.get(key)
            for key in ("smoothing", "baseline")
            if isinstance(processing.get(key), dict)
        }
    signal_payload = _table_payload(signal_pipeline)
    if signal_payload:
        sections.append(("Signal Pipeline", signal_payload))

    analysis_steps = processing.get("analysis_steps")
    if not isinstance(analysis_steps, dict) or not analysis_steps:
        analysis_steps = {
            key: processing.get(key)
            for key in ("glass_transition", "peak_detection", "step_detection")
            if isinstance(processing.get(key), dict)
        }
    step_payload = _table_payload(analysis_steps)
    if step_payload:
        sections.append(("Analysis Steps", step_payload))

    return sections


def _validation_sections(validation: dict | None) -> list[tuple[str, dict[str, str]]]:
    validation = validation or {}
    if not validation:
        return []

    sections: list[tuple[str, dict[str, str]]] = []
    summary = _table_payload(
        {
            "status": validation.get("status"),
            "issue_count": len(validation.get("issues") or []),
            "warning_count": len(validation.get("warnings") or []),
        }
    )
    if summary:
        sections.append(("Data Validation", summary))

    issue_payload = _table_payload({"issues": validation.get("issues")})
    if issue_payload:
        sections.append(("Validation Issues", issue_payload))

    warning_payload = _table_payload({"warnings": validation.get("warnings")})
    if warning_payload:
        sections.append(("Validation Warnings", warning_payload))

    checks_payload = _table_payload(validation.get("checks"))
    if checks_payload:
        sections.append(("Validation Checks", checks_payload))

    return sections


def _provenance_sections(provenance: dict | None) -> list[tuple[str, dict[str, str]]]:
    provenance = provenance or {}
    if not provenance:
        return []

    primary_keys = (
        "saved_at_utc",
        "dataset_key",
        "source_data_hash",
        "vendor",
        "instrument",
        "analyst_name",
        "app_version",
        "analysis_event_count",
    )
    sections: list[tuple[str, dict[str, str]]] = []

    summary = _table_payload({key: provenance.get(key) for key in primary_keys})
    if summary:
        sections.append(("Provenance", summary))

    remaining = {
        key: value
        for key, value in provenance.items()
        if key not in primary_keys and value not in (None, "", [], {})
    }
    context_payload = _table_payload(remaining)
    if context_payload:
        sections.append(("Provenance Context", context_payload))

    return sections


def _record_main_sections(record: dict) -> list[tuple[str, dict[str, Any]]]:
    context_sections = scientific_context_to_report_sections(record.get("scientific_context"))

    methodology_payload: dict[str, Any] = {}
    ordered_sections: list[tuple[str, dict[str, Any]]] = []

    for title, payload in context_sections:
        if title == "Methodology":
            if isinstance(payload, dict):
                methodology_payload.update(payload)
            continue
        if title == "Warnings and Limitations":
            continue
        ordered_sections.append((title, payload if isinstance(payload, dict) else {}))

    domain_summary = _domain_method_summary(record)
    if domain_summary is None:
        domain_summary = _generic_method_summary(record.get("processing"))
    methodology_payload.update(domain_summary or {})

    sections: list[tuple[str, dict[str, Any]]] = []
    if methodology_payload:
        sections.append(("Methodology", _table_payload(methodology_payload)))

    sections.extend(ordered_sections)

    grouped = condense_warning_limitations(
        record.get("scientific_context"),
        validation=record.get("validation"),
        max_data_warnings=4,
        max_method_limits=4,
    )
    warning_payload: dict[str, Any] = {}
    if grouped.get("data_completeness_warnings"):
        warning_payload["Data Completeness Warnings"] = grouped["data_completeness_warnings"]
    if grouped.get("methodological_limitations"):
        warning_payload["Methodological Limitations"] = grouped["methodological_limitations"]
    if warning_payload:
        sections.append(("Warnings and Limitations", warning_payload))
    return sections


def _record_appendix_sections(record: dict) -> list[tuple[str, dict[str, str]]]:
    sections: list[tuple[str, dict[str, str]]] = []
    sections.extend(_processing_sections(record.get("processing")))
    sections.extend(_validation_sections(record.get("validation")))
    sections.extend(_provenance_sections(record.get("provenance")))
    metadata_payload = _appendix_dataset_metadata(record.get("metadata") or {})
    if metadata_payload:
        sections.append(("Dataset Metadata (Technical)", metadata_payload))
    review_payload = _table_payload(record.get("review"))
    if review_payload:
        sections.append(("Internal Review Context", review_payload))
    return sections


def _render_record_mapping(doc: Document, title: str, payload: dict | None) -> None:
    payload = payload or {}
    if not payload:
        return

    doc.add_paragraph(title, style="Heading 3")
    if title == "Scientific Interpretation":
        for value in payload.values():
            doc.add_paragraph(str(value), style="List Bullet")
        doc.add_paragraph()
        return

    if title == "Warnings and Limitations":
        for group, values in payload.items():
            doc.add_paragraph(str(group), style="Heading 4")
            if isinstance(values, list):
                for item in values:
                    doc.add_paragraph(str(item), style="List Bullet")
            else:
                doc.add_paragraph(str(values), style="List Bullet")
        doc.add_paragraph()
        return

    _add_key_value_table(doc, {key: _format_value(value) for key, value in payload.items()})
    doc.add_paragraph()


def _render_main_record_docx(doc: Document, record: dict) -> None:
    doc.add_paragraph(_record_title(record), style="Heading 2")
    key_results = _record_key_results(record)
    if key_results:
        _render_record_mapping(doc, "Key Results", key_results)

    for title, payload in _record_main_sections(record):
        _render_record_mapping(doc, title, payload)

    major_events = _tga_major_events(record)
    if major_events:
        doc.add_paragraph("Major Decomposition Events", style="Heading 3")
        _add_results_table(doc, ["Event", "Midpoint Temperature (°C)", "Mass Loss (%)", "Final Residue (%)"], major_events)
        doc.add_paragraph()
    else:
        compact = _record_compact_rows(record)
        if compact:
            headers, rows = compact
            doc.add_paragraph("Compact Key Table", style="Heading 3")
            _add_results_table(doc, headers, rows)
            doc.add_paragraph()


def _render_appendix_docx(
    doc: Document,
    *,
    records: list[dict],
    datasets: dict,
    comparison_payload: dict[str, Any] | None,
) -> None:
    dataset_sections = []
    for dataset_key, dataset in datasets.items():
        payload = _appendix_dataset_metadata(getattr(dataset, "metadata", {}) or {})
        if payload:
            dataset_sections.append((dataset_key, payload))

    record_sections = []
    for record in records:
        sections = _record_appendix_sections(record)
        full_rows = _record_full_rows(record)
        if sections or full_rows:
            record_sections.append((record, sections, full_rows))

    comparison_has_content = bool(comparison_payload and (comparison_payload.get("appendix_overview") or comparison_payload.get("appendix_batch_rows")))
    if not (dataset_sections or record_sections or comparison_has_content):
        return

    _add_heading(doc, "Appendix A — Reproducibility and Audit Trail", level=1)

    if dataset_sections:
        doc.add_paragraph("Dataset Import and Metadata Technical Details", style="Heading 2")
        for dataset_key, payload in dataset_sections:
            doc.add_paragraph(_dataset_label(dataset_key, datasets), style="Heading 3")
            _add_key_value_table(doc, payload)
            doc.add_paragraph()

    if comparison_has_content:
        doc.add_paragraph("Comparison Workspace Technical Context", style="Heading 2")
        if comparison_payload.get("appendix_overview"):
            _add_key_value_table(doc, comparison_payload["appendix_overview"])
            doc.add_paragraph()
        batch_rows = comparison_payload.get("appendix_batch_rows") or []
        if batch_rows:
            batch_totals = summarize_batch_outcomes(batch_rows)
            _add_key_value_table(
                doc,
                {
                    "Batch Total": batch_totals["total"],
                    "Saved": batch_totals["saved"],
                    "Blocked": batch_totals["blocked"],
                    "Failed": batch_totals["failed"],
                },
            )
            doc.add_paragraph()
            _add_results_table(
                doc,
                ["Run", "Sample", "Template", "Execution", "Validation", "Calibration", "Reference", "Result ID", "Error ID", "Reason"],
                [
                    [
                        _format_value(row.get("dataset_key")),
                        _format_value(row.get("sample_name")),
                        _format_value(row.get("workflow_template")),
                        _format_value(row.get("execution_status")),
                        _format_value(row.get("validation_status")),
                        _format_value(row.get("calibration_state")),
                        _format_value(row.get("reference_state")),
                        _format_value(row.get("result_id")),
                        _format_value(row.get("error_id")),
                        _format_value(row.get("failure_reason")),
                    ]
                    for row in batch_rows
                ],
            )
            doc.add_paragraph()

    for record, sections, full_rows in record_sections:
        doc.add_paragraph(_record_title(record), style="Heading 2")
        for title, payload in sections:
            _render_record_mapping(doc, title, payload)
        if full_rows:
            headers, rows = full_rows
            doc.add_paragraph("Full Raw Data Table", style="Heading 3")
            _add_results_table(doc, headers, rows)
            doc.add_paragraph()


def _add_cover_page(doc: Document, branding: dict | None, license_state: dict | None) -> None:
    branding = branding or {}
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title = branding.get("report_title") or "ThermoAnalyzer Professional Report"
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)

    if branding.get("logo_bytes"):
        try:
            doc.add_picture(io.BytesIO(branding["logo_bytes"]), width=Inches(1.4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_bits = [branding.get("company_name"), branding.get("lab_name")]
    subtitle = " | ".join(bit for bit in subtitle_bits if bit)
    if subtitle:
        subtitle_para.add_run(subtitle)

    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    analyst = branding.get("analyst_name") or "Analyst not specified"
    meta_para.add_run(
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d  %H:%M')} | Analyst: {analyst}"
    )

    if license_state and license_state.get("status") in {"trial", "activated"}:
        license_para = doc.add_paragraph()
        license_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sku = (license_state.get("license") or {}).get("sku", "Professional")
        license_para.add_run(f"License status: {license_state['status']} ({sku})")

    doc.add_page_break()


def generate_docx_report(
    results: dict,
    datasets: dict,
    figures: Optional[dict] = None,
    file_path_or_buffer: Optional[Union[str, io.BytesIO]] = None,
    branding: Optional[dict] = None,
    comparison_workspace: Optional[dict] = None,
    license_state: Optional[dict] = None,
) -> bytes:
    """Generate a DOCX report from normalized stable/experimental records."""
    valid_results, issues = split_valid_results(results)
    stable_results, experimental_results = partition_results_by_status(valid_results)
    all_records = stable_results + experimental_results
    comparison_payload = _build_comparison_payload(comparison_workspace, datasets, all_records)
    executive_rows = _build_executive_summary_rows(all_records, datasets, comparison_payload)

    doc = Document()
    _add_cover_page(doc, branding, license_state)

    _add_heading(doc, "Executive Summary", level=1)
    if executive_rows:
        _add_results_table(
            doc,
            ["Dataset / Set", "Analysis Type", "Key Metrics", "Scientific Interpretation", "Confidence / Limitation"],
            executive_rows,
        )
    else:
        doc.add_paragraph("No analysis results were available for executive summarization.")
    doc.add_paragraph()

    _add_heading(doc, "Experimental Conditions", level=1)
    if not datasets:
        doc.add_paragraph("No dataset metadata available.")
    else:
        for dataset_key, dataset in datasets.items():
            doc.add_paragraph(_dataset_label(dataset_key, datasets), style="Heading 2")
            payload = _main_conditions_payload(dataset_key, dataset)
            if payload:
                _add_key_value_table(doc, payload)
            else:
                doc.add_paragraph("No reader-facing experimental metadata available.")
            doc.add_paragraph()

    if comparison_payload:
        _add_heading(doc, "Comparison Overview", level=1)
        if comparison_payload.get("overview"):
            _add_key_value_table(doc, comparison_payload["overview"])
            doc.add_paragraph()
        if comparison_payload.get("metric_rows"):
            _add_results_table(doc, comparison_payload["metric_headers"], comparison_payload["metric_rows"])
            doc.add_paragraph()
        if comparison_payload.get("interpretation"):
            doc.add_paragraph("Comparison Interpretation", style="Heading 2")
            doc.add_paragraph(str(comparison_payload["interpretation"]))
            doc.add_paragraph()

    _add_heading(doc, "Stable Analyses", level=1)
    if not stable_results:
        doc.add_paragraph("No stable analysis results available.")
    else:
        for record in stable_results:
            _render_main_record_docx(doc, record)

    if experimental_results:
        _add_heading(doc, "Experimental Analyses", level=1)
        doc.add_paragraph("These results are included for reference but remain outside the stable workflow guarantee.")
        for record in experimental_results:
            _render_main_record_docx(doc, record)

    report_notes = (branding or {}).get("report_notes")
    if report_notes:
        _add_heading(doc, "Analyst Notes", level=1)
        doc.add_paragraph(str(report_notes))

    if issues:
        _add_heading(doc, "Skipped Records", level=1)
        for issue in issues:
            doc.add_paragraph(issue, style="List Bullet")

    if figures:
        _add_heading(doc, "Figures", level=1)
        for caption, png_bytes in figures.items():
            doc.add_paragraph(caption, style="Heading 2")
            try:
                img_stream = io.BytesIO(png_bytes)
                doc.add_picture(img_stream, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                doc.add_paragraph("Figure could not be embedded and was skipped.")
            doc.add_paragraph()

    _render_appendix_docx(doc, records=all_records, datasets=datasets, comparison_payload=comparison_payload)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()

    if isinstance(file_path_or_buffer, str):
        with open(file_path_or_buffer, "wb") as fh:
            fh.write(docx_bytes)
    elif isinstance(file_path_or_buffer, io.BytesIO):
        file_path_or_buffer.write(docx_bytes)
        file_path_or_buffer.seek(0)

    return docx_bytes


def _comparison_batch_rows(comparison_workspace: dict | None) -> list[dict]:
    comparison_workspace = comparison_workspace or {}
    return normalize_batch_summary_rows(comparison_workspace.get("batch_summary") or [])


def generate_csv_summary(
    results: dict,
    file_path_or_buffer: Optional[Union[str, io.StringIO]] = None,
) -> str:
    """Generate a flat CSV summary from normalized result records."""
    valid_results, _ = split_valid_results(results)
    flat_rows = flatten_result_records(valid_results)

    fieldnames = [
        "result_id",
        "status",
        "analysis_type",
        "dataset_key",
        "section",
        "row_index",
        "field",
        "value",
    ]

    str_buffer = io.StringIO()
    writer = csv.DictWriter(str_buffer, fieldnames=fieldnames)
    writer.writeheader()
    for row in flat_rows:
        writer.writerow(row)
    csv_str = str_buffer.getvalue()

    if isinstance(file_path_or_buffer, str):
        with open(file_path_or_buffer, "w", newline="", encoding="utf-8") as fh:
            fh.write(csv_str)
    elif isinstance(file_path_or_buffer, io.StringIO):
        file_path_or_buffer.write(csv_str)
        file_path_or_buffer.seek(0)

    return csv_str


def pdf_export_available() -> bool:
    """Return whether reportlab is installed for PDF export."""
    try:  # pragma: no cover - availability depends on environment
        import reportlab  # noqa: F401
    except ImportError:
        return False
    return True


def generate_pdf_report(
    results: dict,
    datasets: dict,
    figures: Optional[dict] = None,
    file_path_or_buffer: Optional[Union[str, io.BytesIO]] = None,
    branding: Optional[dict] = None,
    comparison_workspace: Optional[dict] = None,
    license_state: Optional[dict] = None,
) -> bytes:
    """Generate a narrative PDF report when reportlab is available."""
    try:  # pragma: no cover - optional dependency
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ImportError("PDF export requires reportlab. Install it with: pip install reportlab") from exc

    valid_results, issues = split_valid_results(results)
    stable_results, experimental_results = partition_results_by_status(valid_results)
    all_records = stable_results + experimental_results
    comparison_payload = _build_comparison_payload(comparison_workspace, datasets, all_records)
    executive_rows = _build_executive_summary_rows(all_records, datasets, comparison_payload)

    styles = getSampleStyleSheet()
    story = []

    def add_kv_table(payload: dict[str, Any]) -> None:
        rows = [["Parameter", "Value"]]
        rows.extend([[str(key), _format_value(value)] for key, value in payload.items()])
        table = Table(rows, hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF3F8")]),
                ]
            )
        )
        story.append(table)

    def add_matrix_table(headers: list[str], rows: list[list[Any]]) -> None:
        matrix = [headers] + [[_format_value(value) for value in row] for row in rows]
        table = Table(matrix, hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF3F8")]),
                ]
            )
        )
        story.append(table)

    def append_main_record(record: dict) -> None:
        story.append(Paragraph(_record_title(record), styles["Heading2"]))
        key_results = _record_key_results(record)
        if key_results:
            story.append(Paragraph("Key Results", styles["Heading3"]))
            add_kv_table(key_results)
            story.append(Spacer(1, 0.08 * inch))

        for title, payload in _record_main_sections(record):
            story.append(Paragraph(title, styles["Heading3"]))
            if title == "Scientific Interpretation":
                for value in payload.values():
                    story.append(Paragraph(f"- {value}", styles["Normal"]))
            elif title == "Warnings and Limitations":
                for group, values in payload.items():
                    story.append(Paragraph(str(group), styles["Heading4"]))
                    if isinstance(values, list):
                        for item in values:
                            story.append(Paragraph(f"- {item}", styles["Normal"]))
                    else:
                        story.append(Paragraph(f"- {values}", styles["Normal"]))
            else:
                add_kv_table({str(key): _format_value(value) for key, value in payload.items()})
            story.append(Spacer(1, 0.08 * inch))

        major_events = _tga_major_events(record)
        if major_events:
            story.append(Paragraph("Major Decomposition Events", styles["Heading3"]))
            add_matrix_table(["Event", "Midpoint Temperature (°C)", "Mass Loss (%)", "Final Residue (%)"], major_events)
            story.append(Spacer(1, 0.08 * inch))
        else:
            compact = _record_compact_rows(record)
            if compact:
                headers, rows = compact
                story.append(Paragraph("Compact Key Table", styles["Heading3"]))
                add_matrix_table(headers, rows)
                story.append(Spacer(1, 0.08 * inch))

    title = (branding or {}).get("report_title") or "ThermoAnalyzer Professional Report"
    story.append(Paragraph(title, styles["Title"]))
    story.append(Paragraph(datetime.datetime.now().strftime("Generated: %Y-%m-%d %H:%M"), styles["Normal"]))
    if branding:
        header_bits = [branding.get("company_name"), branding.get("lab_name"), branding.get("analyst_name")]
        header = " | ".join(bit for bit in header_bits if bit)
        if header:
            story.append(Paragraph(header, styles["Normal"]))
    if license_state and license_state.get("status"):
        story.append(Paragraph(f"License: {license_state['status']}", styles["Normal"]))
    story.append(Spacer(1, 0.2 * inch))

    if branding and branding.get("logo_bytes"):
        try:
            story.append(Image(io.BytesIO(branding["logo_bytes"]), width=1.4 * inch, height=0.8 * inch))
            story.append(Spacer(1, 0.2 * inch))
        except Exception:
            pass

    story.append(Paragraph("Executive Summary", styles["Heading1"]))
    if executive_rows:
        add_matrix_table(["Dataset / Set", "Analysis Type", "Key Metrics", "Scientific Interpretation", "Confidence / Limitation"], executive_rows)
    else:
        story.append(Paragraph("No analysis results were available for executive summarization.", styles["Normal"]))
    story.append(Spacer(1, 0.15 * inch))

    story.append(Paragraph("Experimental Conditions", styles["Heading1"]))
    if not datasets:
        story.append(Paragraph("No dataset metadata available.", styles["Normal"]))
    else:
        for dataset_key, dataset in datasets.items():
            story.append(Paragraph(_dataset_label(dataset_key, datasets), styles["Heading2"]))
            payload = _main_conditions_payload(dataset_key, dataset)
            if payload:
                add_kv_table(payload)
            else:
                story.append(Paragraph("No reader-facing experimental metadata available.", styles["Normal"]))
            story.append(Spacer(1, 0.08 * inch))

    if comparison_payload:
        story.append(Paragraph("Comparison Overview", styles["Heading1"]))
        if comparison_payload.get("overview"):
            add_kv_table(comparison_payload["overview"])
            story.append(Spacer(1, 0.08 * inch))
        if comparison_payload.get("metric_rows"):
            add_matrix_table(comparison_payload["metric_headers"], comparison_payload["metric_rows"])
            story.append(Spacer(1, 0.08 * inch))
        if comparison_payload.get("interpretation"):
            story.append(Paragraph("Comparison Interpretation", styles["Heading2"]))
            story.append(Paragraph(str(comparison_payload["interpretation"]), styles["Normal"]))
            story.append(Spacer(1, 0.08 * inch))

    story.append(Paragraph("Stable Analyses", styles["Heading1"]))
    if stable_results:
        for record in stable_results:
            append_main_record(record)
    else:
        story.append(Paragraph("No stable analysis results available.", styles["Normal"]))

    if experimental_results:
        story.append(Paragraph("Experimental Analyses", styles["Heading1"]))
        story.append(Paragraph("These results are included for reference but remain outside the stable workflow guarantee.", styles["Normal"]))
        for record in experimental_results:
            append_main_record(record)

    if (branding or {}).get("report_notes"):
        story.append(Paragraph("Analyst Notes", styles["Heading1"]))
        story.append(Paragraph(str((branding or {})["report_notes"]), styles["Normal"]))

    if issues:
        story.append(Paragraph("Skipped Records", styles["Heading1"]))
        for issue in issues:
            story.append(Paragraph(f"- {issue}", styles["Normal"]))

    if figures:
        story.append(Paragraph("Figures", styles["Heading1"]))
        for caption, png_bytes in figures.items():
            story.append(Paragraph(caption, styles["Heading2"]))
            try:
                story.append(Image(io.BytesIO(png_bytes), width=5.5 * inch, height=3.4 * inch))
            except Exception:
                continue
            story.append(Spacer(1, 0.12 * inch))

    dataset_sections = []
    for dataset_key, dataset in datasets.items():
        payload = _appendix_dataset_metadata(getattr(dataset, "metadata", {}) or {})
        if payload:
            dataset_sections.append((dataset_key, payload))

    record_sections = []
    for record in all_records:
        sections = _record_appendix_sections(record)
        full_rows = _record_full_rows(record)
        if sections or full_rows:
            record_sections.append((record, sections, full_rows))

    comparison_has_content = bool(comparison_payload and (comparison_payload.get("appendix_overview") or comparison_payload.get("appendix_batch_rows")))
    if dataset_sections or record_sections or comparison_has_content:
        story.append(Paragraph("Appendix A - Reproducibility and Audit Trail", styles["Heading1"]))

        if dataset_sections:
            story.append(Paragraph("Dataset Import and Metadata Technical Details", styles["Heading2"]))
            for dataset_key, payload in dataset_sections:
                story.append(Paragraph(_dataset_label(dataset_key, datasets), styles["Heading3"]))
                add_kv_table(payload)
                story.append(Spacer(1, 0.08 * inch))

        if comparison_has_content:
            story.append(Paragraph("Comparison Workspace Technical Context", styles["Heading2"]))
            if comparison_payload.get("appendix_overview"):
                add_kv_table(comparison_payload["appendix_overview"])
                story.append(Spacer(1, 0.08 * inch))
            batch_rows = comparison_payload.get("appendix_batch_rows") or []
            if batch_rows:
                batch_totals = summarize_batch_outcomes(batch_rows)
                add_kv_table(
                    {
                        "Batch Total": batch_totals["total"],
                        "Saved": batch_totals["saved"],
                        "Blocked": batch_totals["blocked"],
                        "Failed": batch_totals["failed"],
                    }
                )
                story.append(Spacer(1, 0.08 * inch))
                add_matrix_table(
                    ["Run", "Sample", "Template", "Execution", "Validation", "Calibration", "Reference", "Result ID", "Error ID", "Reason"],
                    [
                        [
                            _format_value(row.get("dataset_key")),
                            _format_value(row.get("sample_name")),
                            _format_value(row.get("workflow_template")),
                            _format_value(row.get("execution_status")),
                            _format_value(row.get("validation_status")),
                            _format_value(row.get("calibration_state")),
                            _format_value(row.get("reference_state")),
                            _format_value(row.get("result_id")),
                            _format_value(row.get("error_id")),
                            _format_value(row.get("failure_reason")),
                        ]
                        for row in batch_rows
                    ],
                )
                story.append(Spacer(1, 0.08 * inch))

        for record, sections, full_rows in record_sections:
            story.append(Paragraph(_record_title(record), styles["Heading2"]))
            for title, payload in sections:
                story.append(Paragraph(title, styles["Heading3"]))
                add_kv_table(payload)
                story.append(Spacer(1, 0.08 * inch))
            if full_rows:
                headers, rows = full_rows
                story.append(Paragraph("Full Raw Data Table", styles["Heading3"]))
                add_matrix_table(headers, rows)
                story.append(Spacer(1, 0.08 * inch))

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    doc.build(story)
    pdf_bytes = buffer.getvalue()

    if isinstance(file_path_or_buffer, str):
        with open(file_path_or_buffer, "wb") as fh:
            fh.write(pdf_bytes)
    elif isinstance(file_path_or_buffer, io.BytesIO):
        file_path_or_buffer.write(pdf_bytes)
        file_path_or_buffer.seek(0)

    return pdf_bytes
